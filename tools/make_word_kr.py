"""한국어 상세 해설 Word(.docx) — 수식은 이미지(matplotlib mathtext)로 렌더 삽입,
근육명은 영문(한글) 병기 + 용어표. .venv/bin/python tools/make_word_kr.py"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TMP = "/tmp/eq"; os.makedirs(TMP, exist_ok=True)
_n = [0]

# 근육 영문→한글
MUS = {
 "DELT1":"전삼각근","DELT2":"중삼각근","DELT3":"후삼각근","SUPSP":"극상근","INFSP":"극하근",
 "SUBSC":"견갑하근","TMIN":"소원근","TMAJ":"대원근","PECM1":"대흉근(쇄골부)","PECM2":"대흉근(흉골부)",
 "PECM3":"대흉근(복부)","LAT1":"광배근1","LAT2":"광배근2","LAT3":"광배근3","CORB":"오훼완근",
 "TRIlong":"상완삼두근 장두","TRIlat":"상완삼두근 외측두","TRImed":"상완삼두근 내측두","ANC":"주근",
 "SUP":"회외근","BIClong":"상완이두근 장두","BICshort":"상완이두근 단두","BRA":"상완근","BRD":"상완요골근",
 "PT":"원회내근","PQ":"방형회내근","ECRL":"장요측수근신근","ECRB":"단요측수근신근","ECU":"척측수근신근",
 "FCR":"요측수근굴근","FCU":"척측수근굴근","PL":"장장근",
}
def mb(name):  # "DELT2 (중삼각근)"
    return f"{name} ({MUS.get(name,name)})"


def eq(s, fs=19):
    """mathtext 수식 → PNG 경로. mathtext 비호환 토큰 변환."""
    import re
    s2 = (s.replace(r"\tfrac", r"\frac").replace(r"\mathbb{1}", r"\mathbf{1}")
           .replace(r"\mathbb{R}", r"\mathrm{I\!R}"))
    s2 = re.sub(r"\\frac(\d)(\d)", r"\\frac{\1}{\2}", s2)          # \frac12 → \frac{1}{2}
    s2 = re.sub(r"\\text\{[^}]*\}", lambda m: m.group(0) if m.group(0).isascii() else "", s2)  # 수식 내 한글 제거
    _n[0]+=1; p=os.path.join(TMP,f"e{_n[0]}.png")
    fig=plt.figure(figsize=(0.1,0.1))
    fig.text(0.5,0.5,f"${s2}$",fontsize=fs,ha="center",va="center")
    try:
        fig.savefig(p,dpi=200,bbox_inches="tight",pad_inches=0.04)
    except ValueError as e:
        print("EQ FAIL:", s2[:70]); plt.close(fig); raise
    plt.close(fig); return p


def H(d,t,l=1): d.add_heading(t,level=l)
def P(d,t,s=10.5,b=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.size=Pt(s); r.bold=b; return p
def INT(d,t):  # 직관 박스
    p=d.add_paragraph(); r=p.add_run("[직관] "); r.bold=True; r.font.color.rgb=RGBColor(0x1f,0x6f,0xb0)
    r2=p.add_run(t); r2.font.size=Pt(10.5)
def EQ(d,s,fs=19):
    d.add_picture(eq(s,fs)); d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
def tbl(d,headers,rows,fs=9):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; rn=c.paragraphs[0].add_run(h); rn.bold=True; rn.font.size=Pt(fs)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=""; rn=cs[i].paragraphs[0].add_run(str(v)); rn.font.size=Pt(fs)


doc=Document()
doc.add_heading("운동학만으로 근육 약화를 찾는다 — 한국어 상세 해설",0)
sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("쉽고 자세하게 · 수식은 정확하게 · 근육명 영문(한글) 병기 · 2026-06-11"); r.italic=True; r.font.size=Pt(11)

# 1
H(doc,"1. 무엇을·왜 하는가")
P(doc,"가상의 팔(근육 63개 MuJoCo 모형 myoArm)에서 특정 근육의 최대 힘(F_max)과 최적 길이(L_opt)를 "
  "일부러 낮춰 '약화'를 만든다. 약화 정도는 배율 s_F(힘)·s_L(길이)로 적는다(s_F=0.3 → 힘이 30%로 감소). "
  "그 팔이 표준 거상(팔 들어올리기)을 할 때의 관절 움직임(운동학)만 관측해, 어느 근육이 얼마나 약한지를 "
  "되맞히는 회귀기를 만든다. EMG(근전도)는 쓰지 않는다 — 임상에서 모션캡처로 얻는 정보만 쓰는 셈.")
INT(doc,"핵심은 '느슨하게 따라 하기'. 정확히 따라 하게 하면 약한 근육을 다른 근육이 보상해 약화가 안 보인다. "
  "느슨하면 경증=보상(움직임 정상)·중증=목표 미달(이탈)로 갈려 회귀기가 그 차이를 읽는다.")
P(doc,"파이프라인: (1) 사람 데이터로 목표 궤적 생성 → (2) 강화학습으로 '약화를 알고 목표를 느슨히 추종'하는 "
  "정책 학습 → (3) 약화×동작변형 대량 시뮬레이션 → (4) 운동학→약화 회귀기 학습.")

# 2
H(doc,"2. 시뮬레이션 환경")
P(doc,"관절 38개(독립 27+견갑 연동 11), 근육 63개. 제어 50Hz(Δt=0.02s). 행동 근육 26개 = 어깨 15 + 팔꿈치 9 + 회내 2.")
P(doc,"행동 어깨근 15:",b=True)
P(doc,", ".join(mb(m) for m in ["DELT1","DELT2","DELT3","SUPSP","INFSP","SUBSC","TMIN","TMAJ","PECM1","PECM2","PECM3","LAT1","LAT2","LAT3","CORB"]),9.5)
P(doc,"행동 팔꿈치근 9 + 회내 2:",b=True)
P(doc,", ".join(mb(m) for m in ["TRIlong","TRIlat","TRImed","ANC","SUP","BIClong","BICshort","BRA","BRD","PT","PQ"]),9.5)
P(doc,"약화(섭동) 주입 — 근육 i의 MuJoCo gain 파라미터를 직접 수정:")
EQ(doc,r"\mathrm{gain}_{i,2}\leftarrow s_F^{c}\cdot\mathrm{gain}^{0}_{i,2}\ (=F_{max}),\quad \mathrm{gain}_{i,0:2}\leftarrow \mathrm{gain}^{0}_{i,0:2}/s_L^{c}")
P(doc,"섭동 채널 10개 (개별 8 + 묶음 2):",b=True)
tbl(doc,["채널","근육(영문/한글)"],[
  ["DELT1",mb("DELT1")],["DELT2",mb("DELT2")],["DELT3",mb("DELT3")],["SUPSP",mb("SUPSP")],
  ["PECM1",mb("PECM1")],["CORB",mb("CORB")],["BIClong",mb("BIClong")],["TRIlong",mb("TRIlong")],
  ["A_lowcuff","하부 회전근개: "+", ".join(mb(m) for m in ["INFSP","SUBSC","TMIN"])],
  ["B_latadd","광배·대원근: "+", ".join(mb(m) for m in ["TMAJ","LAT1","LAT2","LAT3"])]])
P(doc,"희소-k 커리큘럼: 매 에피소드 손상 채널 수 k∈{1,2,3}(가중 0.5/0.3/0.2)만 약화, 그 채널만 s_F∼U[0.05,1], s_L∼U[0.7,1.2].",9.5)

# 3 RL
H(doc,"3. 강화학습 정책 (PPO)")
H(doc,"3.1 관측 (정책 입력, 76차원)",2)
INT(doc,"매 순간 '지금 자세·속도, 목표와 오차, 이 동작 특성과 어느 근육이 약한지'를 한 벡터로 준다.")
P(doc,"관측은 상태의 함수 o_t ∈ ℝ^76. 추적 자유도 D={거상, 평면, 팔꿈치, 회내}, 자유축=축회전.")
EQ(doc,r"\mathbf{o}_t=\left[\ \tfrac{1}{\pi}\mathbf{q}\,(5),\ \tfrac{1}{5}\dot{\mathbf{q}}\,(5),\ \mathbf{a}^{act}\,(26),\ \tfrac{1}{\pi}\mathbf{q}^{ref}\,(4),\ \tfrac{1}{\pi}\mathbf{e}\,(4),\ \tfrac{1}{\pi}\mathbf{q}^{ref}_{+5}\,(4),\ \varphi,\mathbb{1}_{hold}\,(2),\ \tilde{\ell}\,(5),\ s_F,s_L\,(20),\ \mathrm{task}\,(1)\right]",16)
tbl(doc,["#","블록","차원","의미"],[
  ["1,2","관절각 q, 속도 q̇","5,5","현재 자세·운동상태 (÷π, ÷5)"],
  ["3","근육 활성 a^act","26","자기 근육 활성도 [0,1] (협응·떨림제어)"],
  ["4,5,6","참조·오차 e·선행","4,4,4","지금 목표, 얼마나 벗어났나, 0.1s 뒤 목표"],
  ["7","위상 φ, hold","2","동작 진행도(t/H), 정점유지 여부"],
  ["8","latent ℓ̃","5","동작 변동인자(속도T·정점·모양·평면·팔꿈치)"],
  ["9","섭동 (s_F,s_L)","20","약화 정보(채널10×2). 만성환자=적응 가정"],
  ["10","과제","1","T1(관상)/T2(전방)"]])
P(doc,"합 5+5+26+4+4+4+2+5+20+1=76. 신경망 입력 전 실행통계로 표준화(clip ±10).",9.5)

H(doc,"3.2 행동 (정책 출력) + 운동노이즈",2)
EQ(doc,r"\mathcal{A}=[-1,1]^{26},\qquad u_j=\tfrac12(a_j+1)\in[0,1]\ (\text{근육 흥분})")
P(doc,"신호의존 운동노이즈(Harris–Wolpert 1998; 노이즈 크기가 명령에 비례 → 생리적 변동):")
EQ(doc,r"u_j\leftarrow \mathrm{clip}(u_j+\eta_j,0,1),\quad \eta_j\sim\mathcal{N}\!\left(0,(\sigma_{sd}u_j+\sigma_0)^2\right),\ \ \sigma_{sd}=0.1,\ \sigma_0=0.01")

H(doc,"3.3 보상 — '잘했다'의 정의 (항별)",2)
EQ(doc,r"r_t=\mathrm{TASK}_t\cdot\mathrm{QUALITY}_t+w_e\,\mathrm{EFFORT}_t+\mathrm{pen}^{rot}_t",20)
P(doc,"(1) TASK — 양면 느슨밴드:")
EQ(doc,r"\mathrm{TASK}_t=\sum_{i\in D} w_i\,\exp\!\left(-k_i^{out}\left[\max(0,\,|q_i-q^{ref}_i|-b_i)\right]^2\right)")
INT(doc,"목표각에서 밴드 b_i(약 10–15°) 이내면 벌점 0(사람마다 다른 건 허용), 밖이면 급감. 거상 가중을 0.15로 "
  "낮춰 느슨하게 둬야 중증 약화가 '이탈'로 드러난다. 계수 w=(0.15,0.15,0.25,0.20), b=(10,15,12,20)°.")
P(doc,"(2) QUALITY — 부드러움(가속·행동변화·정점정지 벌점):")
EQ(doc,r"\mathrm{QUALITY}_t=\exp\!\left(-\left(c_J\,\overline{\ddot q^2}+c_D\,\overline{(\mathbf{a}_t-\mathbf{a}_{t-1})^2}+c_S\,\mathbb{1}_{hold}\,\overline{\dot q^2}\right)\right)")
P(doc,"(3) EFFORT(힘 아끼기) + 회전 정칙화:")
EQ(doc,r"\mathrm{EFFORT}_t=-\tfrac{1}{26}\sum_j (a^{act}_j)^2,\qquad \mathrm{pen}^{rot}_t=-w_r\left[\max(0,|q_{rot}|-50^{\circ})\right]^2")
P(doc,"추적오차 50°(거상)/60°(팔꿈치) 초과 또는 속도폭주 시 종료(보상 −1). 밴드를 넓게 둬 약화 이탈은 살리고 발산만 차단.",9.5)

H(doc,"3.4 정책·가치망 출력공간과 PPO",2)
INT(doc,"정책은 '관측→행동의 확률분포'를, 가치망은 '이 상태가 앞으로 얼마나 좋은지'를 낸다.")
EQ(doc,r"\pi_\psi(\cdot|\mathbf{o})=\mathcal{N}\!\left(\mu_\psi(\mathbf{o}),\,\mathrm{diag}(e^{2\sigma_\psi})\right),\qquad V_\psi:\mathcal{O}\to\mathbb{R}")
P(doc,"정책 출력공간 = ℝ^26 × ℝ_{>0}^26 (평균 μ_ψ(o), log-std σ_ψ). μ_ψ·V_ψ는 각각 MLP[256,256]+tanh.")
EQ(doc,r"\mathcal{L}(\psi)=-\mathbb{E}\!\left[\min(\rho_t\hat A_t,\ \mathrm{clip}(\rho_t,1{-}\epsilon,1{+}\epsilon)\hat A_t)\right]+c_v\mathbb{E}[(V_\psi-\hat R_t)^2]-c_e\mathbb{E}[\mathcal{H}(\pi_\psi)]",16)
P(doc,"ε=0.2, lr=3e-4, rollout 512×20, minibatch 4096, 10 epoch, γ=0.99, λ=0.95. 커리큘럼: M1 건강·고정 → M2 분포+T2 → M3 섭동 k=1→k≤3 (warm-start).",9.5)

# 4 reference
H(doc,"4. 참조(따라 할 목표) 생성")
INT(doc,"한 동작만 쓰면 단조롭다. 사람 데이터(KIMHu 20명) 거상 궤적을 다양·현실적으로 변형. 방법: warp(평균+1모드)→VAE→fPCA(채택).")
P(doc,"데이터셋: 반복별 다채널 궤적을 위상 N점 리샘플. X^(m) ∈ ℝ^{N×d}, (N,d)=(50,4) 또는 (80,6), M≈200. "
  "채널=거상·팔꿈치·평면·회내(+손목 굴곡·편위).")
H(doc,"4.1 fPCA / ProMP (채택)",2)
P(doc,"① 평균 x̄, 중심화 후 SVD로 주성분 φ_k·고윳값 λ_k. ② 95% 분산 모드 수 K(실측 13~24). "
  "③ 점수 표준화. ④ 생성=점수공간 KDE(실제 점수 부트스트랩+jitter):")
EQ(doc,r"\hat{z}=\hat{w}^{(j)}+\varepsilon,\ \varepsilon\sim\mathcal{N}(0,h^2 I),\ h{=}0.35;\quad \hat{x}=\bar{x}+\sum_{k=1}^{K}\hat z_k\sqrt{\lambda_k}\,\phi_k")
INT(doc,"왜 좋은가: z∼N(0,I)면 생성 공분산=Σλ_k φ_k φ_k^T=실제 데이터 공분산이라 분포 '폭'이 실제와 같다. "
  "실측 정점 실제 94±9° vs fPCA 95±9°(폭 일치), 베끼지 않음. VAE는 폭이 좁았다(평균화).")
H(doc,"4.2 VAE (비교군)",2)
EQ(doc,r"\mathcal{L}_{VAE}=\|x-\hat x\|_2^2+\beta\cdot\tfrac12\sum_k(\sigma_{z,k}^2+\mu_{z,k}^2-1-\log\sigma_{z,k}^2),\ \ \beta{=}0.5")
P(doc,"인코더 MLP(200→128→128→z6), 디코더(6→128→128→200), 재매개화 z=μ+σ⊙ε. 한계: KL+소량데이터로 분포 과소(정점 sd 9→4°), 잠재 붕괴(6중 3). → fPCA 채택.",9.5)

# 5 regressor
H(doc,"5. 식별 회귀기 (시계열 TCN)")
EQ(doc,r"g:\ \mathbb{R}^{T\times C}\ (\text{운동학})\times \mathbb{R}^{6}\ (\text{공변량})\ \to\ \mathbb{R}^{2|\mathcal{C}|}\ (\hat s_F,\hat s_L)")
P(doc,"입력 = 운동학 시계열만(추적DoF+회전 각도·속도 + 추적오차; C=14 기본/20 손목). 근육활성·EMG 안 씀. 가변길이 제로패딩+마스크.")
P(doc,"구조 = dilated 1D-TCN(시간 다운샘플 conv → dilation 2,4,8) + 마스크 평균·최대 풀링 + 공변량 결합 → MLP. "
  "손실 = 중증 재가중 + Huber. Adam+cosine, 100~120 epoch.")
INT(doc,"왜 시계열인가: 같은 운동학을 요약통계로 회귀하면 R²≈0.36에 그쳤지만, 시계열로 동역학(언제·어떻게 이탈하는지)을 살리면 0.81로 급등.")

# 6 results
H(doc,"6. 결과와 해석")
P(doc,"운동학만(EMG 없음), fPCA 참조, 운동노이즈 하 s_F(힘 배율) 복원:")
tbl(doc,["설정","입력채널","어깨/팔꿈치 R²","손목근 R²"],[
  ["거상(rise), T1+T2, 손목잠금","14","0.81","—"],
  ["전체동작, 손목해제, T1","20","0.60","0.81"],
  ["전체동작, 손목해제, T1+T2","20","0.55","0.74"]])
P(doc,"채널별(거상·최고설정 s_F R²): "
  + f"{mb('SUPSP')} .87, {mb('BIClong')} .87, {mb('TRIlong')} .86, {mb('DELT1')} .84, "
  + f"{mb('DELT2')} .79, {mb('DELT3')} .75, {mb('CORB')}/{mb('PECM1')} .62–.65, "
  + "하부회전근개(A) .94, 광배·대원근(B) .88.",9.5)
P(doc,"해석:",b=True)
for t in ["(1) EMG 없이 움직임만으로 어깨/팔꿈치 근육 힘 약화를 강식별(R²=0.81).",
          "(2) 시계열이 핵심(요약통계 0.36 → 시계열 0.81). 약화는 동작의 시간 구조에 부호화됨.",
          "(3) 참조 생성은 fPCA가 최적(실제 분포 충실, VAE 과소분산 회피).",
          "(4) 보상의 느슨밴드가 약화를 신호로 바꾸는 장치(경증=보상, 중증=이탈).",
          "(5) 트레이드오프: 손목해제+전체동작은 손목근 식별(~0.8) 추가하나, 정보 적은 구간이 거상-rise(어깨근 최정보)를 희석해 어깨/팔꿈치 식별 저하. 한 동작으로 둘 다 최대화 불가."]:
    P(doc,t,10)

out=os.path.join(ROOT,"보고서_한국어상세_2026-06-11.docx")
doc.save(out)
print("saved:",out,"| 문단",len(doc.paragraphs),"표",len(doc.tables),"수식이미지",len(doc.inline_shapes))
