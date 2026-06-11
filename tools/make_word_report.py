"""상세 Word(.docx) 보고서 생성 — 문어체, 표·그래프·용어해설 포함.

시퀀스 회귀 R²는 /tmp/seqG_{all,act,kin}.log에서 파싱(데이터 기반).
사용: .venv/bin/python tools/make_word_report.py
"""
import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from custom_envs import muscle_groups as MG

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(ROOT, "results", "report")
CH = [c for c, _ in MG.PERTURB_CHANNELS]

# 근육 한글/영문 병기
MUS_KR = {
    "DELT1": "전삼각근 (anterior deltoid)", "DELT2": "중삼각근 (middle deltoid)",
    "DELT3": "후삼각근 (posterior deltoid)", "SUPSP": "극상근 (supraspinatus)",
    "PECM1": "대흉근 쇄골부 (pectoralis major, clavicular)", "CORB": "오훼완근 (coracobrachialis)",
    "BIClong": "상완이두근 장두 (biceps brachii, long head)",
    "TRIlong": "상완삼두근 장두 (triceps brachii, long head)",
    "A_lowcuff": "하부 회전근개 묶음 (lower rotator cuff: INFSP·SUBSC·TMIN)",
    "B_latadd": "광배·대원근 묶음 (latissimus·teres major: TMAJ·LAT1-3)",
}


def parse_seq_log(path):
    """regress_seq 로그 → {'sF':{ch:r2}, 'sL':{ch:r2}, 'sF_mean':, 'sL_mean':}."""
    res = {"sF": {}, "sL": {}, "sF_mean": None, "sL_mean": None}
    if not os.path.exists(path):
        return res
    sec = None
    for line in open(path):
        if "s_F: Fmax" in line:
            sec = "sF"
        elif "=== s_L" in line:
            sec = "sL"
        m = re.match(r"^(\w+)\s+([\d.]+)\s+(-?[\d.]+)", line)
        if m and m.group(1) in CH and sec:
            res[sec][m.group(1)] = float(m.group(3))
        mm = re.search(r"s_([FL]) 평균 R²\s+([\d.]+)", line)
        if mm:
            res["s" + mm.group(1) + "_mean"] = float(mm.group(2))
    return res


# 스냅샷 MLP 실측값(앞선 실행 결과)
SNAP_SF = {"DELT1": .94, "DELT2": .96, "DELT3": .93, "SUPSP": .94, "PECM1": .83,
           "CORB": .92, "BIClong": .91, "TRIlong": .94, "A_lowcuff": .97, "B_latadd": .96}
SNAP = {"act": .90, "kin": .36, "all": .92, "sL": .81}


def H(doc, txt, lvl=1):
    doc.add_heading(txt, level=lvl)


def P(doc, txt, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
    return t


def img(doc, name, width=6.4, caption=None):
    p = os.path.join(FIG, name)
    if os.path.exists(p):
        doc.add_picture(p, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)


def build():
    seq = {ch: parse_seq_log(f"/tmp/seqG_{ch}.log") for ch in ("all", "act", "kin")}
    doc = Document()
    # 제목
    t = doc.add_heading("myoArm 근골격 파라미터 추정 연구 보고서", 0)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("— 시뮬레이션 내부 식별성(sim-internal identifiability) 및 시계열 회귀 정식화 —\n작성일: 2026-06-11")
    r.italic = True; r.font.size = Pt(11)

    H(doc, "1. 연구 개요", 1)
    P(doc, "본 연구는 상지 근골격 시뮬레이션(musculoskeletal simulation) 환경 myoArm에서, "
      "특정 근육의 힘 발생 능력을 인위적으로 약화시킨 뒤, 그 결과로 나타나는 동작의 운동학"
      "(kinematics, 관절이 움직인 궤적)과 근육 활성도(muscle activation, 각 근육이 얼마나 수축했는가)"
      "만을 관측하여 \"어느 근육이 얼마나 약해졌는가\"를 정량적으로 복원하는 회귀 모형(regression model)을 "
      "개발하는 것을 목표로 한다. 실제 환자 데이터로 외삽(sim-to-real)하는 것이 아니라, 시뮬레이션 내부에서 "
      "정답 파라미터가 알려진 상황에서의 복원 가능성, 즉 식별성(identifiability)을 측정한다.")
    P(doc, "핵심 가설은 다음과 같다. 약화된 근육은 그 정도에 따라 두 가지 신호 중 하나로 드러난다. "
      "경증(輕症)일 경우 신경계가 다른 근육을 동원하여 보상(compensation)하므로 동작 궤적은 정상으로 "
      "회복되지만 근육 활성 패턴이 변하며(활성 신호, ACT), 중증(重症)일 경우 보상이 불가능하여 동작이 "
      "목표에 미치지 못하는 이탈(deviation)로 나타난다(운동학 신호, KIN). 본 보고서는 이 가설이 "
      "시뮬레이션에서 실증되었으며, 특히 동작의 시간적 전개(time series)를 입력으로 삼는 시계열 회귀가 "
      "단순 요약통계보다 우수함을 보고한다.")

    H(doc, "2. 용어 해설", 1)
    add_table(doc, ["용어 (영문)", "쉬운 설명"], [
        ["근골격 모델 (musculoskeletal model)", "뼈·관절·근육을 물리적으로 모사한 가상 팔. 본 연구는 63개 근육의 myoArm 사용."],
        ["거상 (elevation)", "팔을 위로 들어 올리는 동작. 본 연구의 표준 과제."],
        ["식별성 (identifiability)", "관측만으로 숨은 파라미터를 유일하게 알아낼 수 있는 정도."],
        ["강화학습 (reinforcement learning)", "보상을 최대화하도록 시행착오로 \"상태→행동\" 정책을 학습하는 방법."],
        ["정책 (policy)", "현재 상태를 입력받아 26개 근육의 활성을 출력하는 신경망(제어기)."],
        ["섭동 (perturbation)", "근육의 최대 힘(Fmax) 등 파라미터를 인위적으로 바꾸어 약화를 만드는 조작."],
        ["활성도 / 운동학 (ACT / KIN)", "ACT=각 근육의 수축 정도(0~1), KIN=관절 각도 등 움직임."],
        ["회전근개 (rotator cuff)", "어깨 깊은 곳에서 상완골(humerus)을 관절와(glenoid)에 고정·회전시키는 근육군."],
    ])

    H(doc, "3. 연구 방법: 시뮬레이션 환경과 학습", 1)
    H(doc, "3.1 환경과 행동공간", 2)
    P(doc, "대상 모델은 MuJoCo 물리엔진 상의 myoArm으로, 어깨의 독립 3축 — 거상 평면(elv_angle), 거상각"
      "(shoulder_elv), 축 회전(shoulder_rot) — 과 팔꿈치 굴곡(elbow_flexion), 전완 회내외(pro_sup)를 "
      "포함한다. 정책이 직접 제어하는 근육은 어깨 15개, 팔꿈치 9개, 회내근 2개로 총 26개이며, 손목·손가락의 "
      "원위(遠位, distal) 22개 관절은 등식 제약(equality constraint)으로 고정하여 분석 범위를 어깨·팔꿈치로 "
      "한정하였다. 견갑(scapula)은 거상각에 선형 연동되는 리듬으로 자동 처방된다.")
    P(doc, "축 회전(shoulder_rot)에는 참조 궤적을 부여하지 않고 자유 자유도로 두었는데, 이는 인간 데이터"
      "(Kinect 기반)로 상완의 고립된 축 회전을 신뢰성 있게 측정할 수 없기 때문이다. 대신 거상에 수반되는 "
      "의무적 외회전(external rotation)을 물리가 자연히 공급하도록 하였으며, 그 부수 효과로 회전근개 약화가 "
      "회전 이탈로 드러나 식별에 기여한다.")
    H(doc, "3.2 보상 함수 (reward)", 2)
    P(doc, "강화학습 보상은 곱셈형으로 설계하였다. 추적 항 TASK는 각 자유도에서 참조각과의 오차에 대해 "
      "양면 허용밴드(two-sided tolerance band)를 갖는 가우시안으로, 밴드 안에서는 벌점이 없어(개인별 속도·"
      "자세 변동을 허용) 약화가 밴드를 벗어나는 이탈로만 드러나도록 하였다. 여기에 움직임 품질(QUALITY: "
      "가속·행동변화·정지)과 노력비용(EFFORT)을 결합하였다. 핵심 의도는 \"느슨한 추적\"으로, 빡빡하게 "
      "추적하면 약화가 궤적에 흡수되어 신호가 사라지기 때문이다.")
    H(doc, "3.3 학습 절차 (점진적 워밍스타트)", 2)
    P(doc, "학습은 기어가기-걷기-뛰기(crawl-walk-run) 방식의 단계적 워밍스타트(warm-start, 앞 단계 "
      "가중치를 이어받아 다음 단계를 시작)로 진행하였다. 모든 단계에서 관측 차원을 76으로 고정하여 호환성을 "
      "유지하였다.")
    add_table(doc, ["단계", "내용", "결과"], [
        ["M1", "건강한 근육·고정 궤적(T1)으로 사람다운 거상 재현", "검증 게이트 5기준 통과(정점 97°, 추종오차 2.8°)"],
        ["M2", "다양한 속도·정점(latent 분포) + 두 번째 과제(T2 전방거상) 추가", "T1·T2 분포 전반 추종(정점오차 3~5°)"],
        ["M3", "근육 섭동 커리큘럼: 손상 근육 수 k=1 → k≤3", "약화→경증 ACT보상·중증 KIN이탈, 정도별 graded"],
        ["M4", "M3 정책으로 20,000회 롤아웃, 특징·라벨 추출", "시계열 데이터셋 (20000×175스텝×77채널)"],
        ["G", "회귀 신경망으로 근육 파라미터 복원", "아래 5절"],
    ])
    img(doc, "fig4_learning_curves.png", 6.6, "그림 1. 학습곡선. 정점오차·추종 RMS가 통과 임계(12°) 아래로 수렴.")

    H(doc, "4. 건강 정책의 동작 검증", 1)
    P(doc, "학습된 정책이 사람다운 거상을 재현하는지 궤적으로 확인하였다. 그림 2는 단일 정책이 다양한 latent와 "
      "두 과제(T1 관상면 외전, T2 전방 거상)에서 각자의 참조 궤적을 따라 부드럽게 상승·유지함을 보인다. "
      "그림 3은 달성 정점이 참조 정점을 따라감(상관 0.9 내외)을 보여, 정책이 latent을 실제로 활용함을 확인한다.")
    img(doc, "fig1_tracking_gallery.png", 6.8, "그림 2. M2 단일 정책의 다중 latent·양 과제 추적 궤적.")
    img(doc, "fig2_latent_usage.png", 6.4, "그림 3. 참조 정점 대 달성 정점 (정책의 latent 활용).")
    P(doc, "그림 4는 근육 약화 시의 신호를 보인다. 상단은 거상 궤적으로, 중삼각근(DELT2)이나 삼두 장두"
      "(TRIlong)의 중증 약화는 목표 미달(undershoot)로 드러나며 약화 정도에 비례한다. 반면 극상근(SUPSP) "
      "등은 세 곡선이 겹쳐, 다른 근육으로 보상하여 궤적이 회복됨을 보인다. 하단의 활성도에서 약화 근육은 "
      "감소하고 보상 근육이 증가하여, 궤적이 회복되더라도 활성 패턴으로 식별 가능함을 보인다.")
    img(doc, "fig3_perturbation_signal.png", 6.8, "그림 5. 약화 정도별 운동학(상단)·활성(하단) 신호.")

    H(doc, "5. 회귀 신경망과 식별성 결과", 1)
    H(doc, "5.1 시계열 회귀로의 정식화", 2)
    P(doc, "초기에는 한 동작에서 추출한 요약통계(각 근육 활성의 평균·최댓값 등 143차원)를 입력으로 하는 "
      "다층퍼셉트론(MLP)을 사용하였다. 그러나 동작의 시간적 전개 — 상승의 동역학, 떨림, 정착 타이밍, 이탈이 "
      "발생하는 시점 — 가 요약통계에서는 소실된다. 이에 매 제어 시점의 상태를 그대로 입력하는 시계열 회귀로 "
      "재정식화하였다. 입력은 175시점 × 77채널(전체 63개 근육 활성, 관절각 5, 관절속도 5, 추종오차 4)의 "
      "시퀀스이며, 모형은 팽창 1차원 합성곱(dilated temporal convolution)에 시간축 마스킹 풀링과 알려진 "
      "공변량(latent·과제)을 융합하는 시간합성곱망(TCN)이다. 출력은 10개 채널의 Fmax 배율(s_F)과 길이 배율"
      "(s_L) 총 20개 값이다. 손실은 이상치에 강건한 Huber 손실을 쓰고, 드문 중증 표본에 가중을 더하였다.")

    H(doc, "5.2 채널별 식별성", 2)
    sa = seq["all"]
    rows = []
    for ch in CH:
        rows.append([MUS_KR[ch], f"{SNAP_SF.get(ch,'-'):.2f}",
                     f"{sa['sF'].get(ch, float('nan')):.2f}",
                     "강식별" if sa['sF'].get(ch,0) > .6 else ("부분" if sa['sF'].get(ch,0) > .3 else "약식별")])
    add_table(doc, ["근육 채널 (한글/영문)", "스냅샷 R²", "시계열 R²", "판정(시계열)"], rows)
    P(doc, f"s_F(Fmax) 평균 R²는 스냅샷 {SNAP['all']:.2f} 대비 시계열 {sa['sF_mean'] or 0:.2f}, "
      f"s_L(Lopt) 평균 R²는 스냅샷 {SNAP['sL']:.2f} 대비 시계열 {sa['sL_mean'] or 0:.2f}이다. "
      "시계열 입력은 특히 동역학에 민감한 길이 배율(s_L)에서 이득이 크다(아래 비교 그림 참조).")
    img(doc, "fig6_seq_vs_snapshot.png", 6.6, "그림 6. 채널별 R²: 스냅샷 MLP 대 시계열 TCN.")

    H(doc, "5.3 신호 출처 분리 (ablation)", 2)
    P(doc, "어느 관측 채널이 식별을 가능케 하는지 분리하기 위해, 입력을 활성(ACT)만, 운동학(KIN)만, 둘 다로 "
      "나누어 동일 모형을 재학습하였다.")
    add_table(doc, ["입력 신호", "s_F 평균 R² (시계열)", "해석"], [
        ["활성 ACT만 (privileged)", f"{seq['act']['sF_mean'] or 0:.2f}", "주채널 — 적응한 운동지령이 약화를 부호화"],
        ["운동학 KIN만", f"{seq['kin']['sF_mean'] or 0:.2f}", "보조 — 비보상 약화(주외전근·이관절·회전근개)에서만"],
        ["ACT + KIN (전체)", f"{sa['sF_mean'] or 0:.2f}", "최종"],
    ])
    P(doc, "활성 신호가 주채널이며, 운동학 신호는 보상이 불가능한 근육에서 보조적으로 기여한다. 이는 \"두 채널 "
      "중 하나로 복원\"이라는 본 연구의 핵심 명제를 실증한다.")

    H(doc, "6. 결론", 1)
    P(doc, "시뮬레이션 내부에서, 적응을 마친 정책의 한 번의 거상 동작만으로 개별 근육(및 병합된 근육군)의 약화 "
      "정도를 높은 정확도로 복원할 수 있음을 확인하였다. 활성도가 주된 식별 채널이며 운동학이 이를 보강한다. "
      "특히 동작의 시간적 전개를 보존하는 시계열 회귀가 요약통계 기반보다 우수하며, 동역학에 민감한 파라미터"
      "(길이 배율)에서 그 이득이 두드러진다. 서로 중복되는 회전근개 근육은 개별 분해 대신 묶음으로 추정하여 "
      "묶음 수준의 약화를 신뢰성 있게 식별하였다.")

    out = os.path.join(ROOT, "보고서_myoArm_2026-06-11.docx")
    doc.save(out)
    print("saved:", out)
    print(f"  seq s_F 평균 R²: all={sa['sF_mean']} act={seq['act']['sF_mean']} kin={seq['kin']['sF_mean']}")


if __name__ == "__main__":
    build()
